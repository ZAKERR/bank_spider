# -*- coding: utf-8 -*-
import json
import re
from functools import reduce

import requests
from fake_useragent import UserAgent


def parse_docx(docx_path):
    """ 直接读取docx """
    from docx import Document
    li = []
    re_com = re.compile(r'\r|\n|\t|\s')
    # 打开文档
    document = Document(docx_path)
    tables = document.tables  # 获取文件中的表格集
    table = tables[0]  # 获取文件中的第一个表格
    for i in range(1, len(table.rows)):  # 从表格第二行开始循环读取表格数据
        result = table.cell(i, 0).text + "" + table.cell(i, 1).text + table.cell(i, 2).text + table.cell(i, 3).text
        li.append(result)

    ws_nr_txt = reduce(lambda x, y: x + y, [re_com.sub('', i) for i in li])
    return ws_nr_txt
    # # 处理纯文本
    # docx_text = [paragraph.text for paragraph in document.paragraphs]
    # return docx_text


def handle_text(txt):
    headers = {'User-Agent': UserAgent().random}
    url = 'http://117.78.35.12:9613/wsfx?txt={}'.format(txt)
    try:
        res = requests.post(url=url, headers=headers)
        if res.status_code == 200:
            return json.loads(res.text)
        return None
    except Exception as e:
        handle_text(txt)
        print(e.args)


if __name__ == '__main__':
    path = r'C:\xiugai\bank_spider\files\word\黄冈银保监分局行政处罚信息公开表（黄冈银保监罚决字〔2019〕6号）.docx'
    txt = parse_docx(path)
    res = handle_text(txt)
    print(res)